from django.contrib.auth import authenticate, login, get_user_model
from django.http import HttpResponse
from django.shortcuts import redirect, render
from django.views import View
from django.utils.http import urlsafe_base64_decode
from django.core.exceptions import ValidationError
from django.contrib.auth.tokens import default_token_generator as token_generator
from django.contrib.auth.views import LoginView
from users.forms import UserCreationForm, AuthenticationForm, UserDataForm
from django.contrib.auth.decorators import login_required
from users.utils import send_email_for_verify
from django import forms
from django.urls import reverse
from django.views.decorators.http import require_GET
from django.http import JsonResponse
from datetime import time
from datetime import datetime
import json
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
from docxtpl import DocxTemplate
from django.http import FileResponse
from io import BytesIO
import io
from django.contrib import messages
from docx import Document
import os
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.contrib.auth.models import Group
from django.contrib.auth import get_user_model
from django.utils.http import quote
from django.db.models import Q
from django.db.models import Count
import calendar
from django.db.models import Sum
from django.utils import timezone
from datetime import timedelta
from django.core.mail import send_mail
from django.db import IntegrityError

from .models import Instructor
from .models import Appointment, Instructor, Notation, UserData, Payments, Message
from .forms import AppointmentForm, UserPaymentsForm
from django.shortcuts import get_object_or_404



User = get_user_model()

class MyLoginView(LoginView):
    form_class = AuthenticationForm

class EmailVerify(View):
    
    def get(self, request, uidb64, token):
        user = self.get_user(uidb64)

        if user is not None and token_generator.check_token(user, token):
            user.email_verify = True
            user.save()
            login(request, user)

            if user.groups.filter(name='Преподаватель').exists():
                return redirect('teacher_dashboard')
            elif user.groups.filter(name='Ученик').exists():
                return redirect('student_dashboard')
            if user.groups.filter(name='Менеджер').exists():
                return redirect('manager_dashboard')
        return redirect('invalid_verify')

    @staticmethod
    def get_user(uidb64):
        try:
            uid = urlsafe_base64_decode(uidb64).decode()
            user = User.objects.get(pk=uid)
        except (TypeError, ValueError, OverflowError,
                User.DoesNotExist, ValidationError):
            user = None
        return user


class Register(View):

    template_name = 'registration/register.html'

    def get(self, request):
        context = {
            'form': UserCreationForm()
        }
        return render(request, self.template_name, context)

    def post(self, request):
        form = UserCreationForm(request.POST)

        if form.is_valid():
            # try:
            form.save()
            email = form.cleaned_data.get('email')
            password = form.cleaned_data.get('password1')
            user = authenticate(email=email, password=password)
            send_email_for_verify(request, user)
            return redirect('confirm_email')
            # except IntegrityError:
            #     form.add_error('username', 'Такой пользователь уже существует.')
        else:
            print(form.errors)

        context = {
            'form': form 
        }
        return render(request, self.template_name, context)
    

def user_profile(request):
    if request.user.is_authenticated:
        if request.user.groups.filter(name='Преподаватель').exists():
            return redirect('teacher_dashboard/')
        elif request.user.groups.filter(name='Ученик').exists():
            return redirect('student_dashboard/')
        elif request.user.groups.filter(name='Менеджер').exists():
            return redirect('manager_dashboard/')
    else:
        return redirect('login')  
    

def student_dashboard(request):
    items = Instructor.objects.all()
    appointments = Appointment.objects.filter(date__gte=datetime.now().date())
    
    appointments_data = []
    for appointment in appointments:
        appointments_data.append({
            'id': appointment.id,
            'date': appointment.date.strftime('%d'),
            'time': appointment.time.strftime('%H:%M:%S'),
            'instructor': appointment.instructor.id,
        })

    appointments_json = json.dumps(appointments_data)
    context = {
        'items':items,
        'appointments':appointments_json
    }
    return render(request, "student_dashboard.html", context)

def student_pickDataTime(request):
    items = Instructor.objects.all()
    context = {
        'items':items
    }
    return render(request, "student_pickDataTime.html", context)


def indexInstructor(request, idInstructor):
    item = Instructor.objects.get(id=idInstructor)
    
    return redirect('appointment', idInstructor=idInstructor)


available_times = [
    time(8, 0),
    time(9, 30),
    time(11, 0),
    time(13, 30),
    time(15, 0),
    time(16, 30),
]

def appointment(request, idInstructor):
    instructor = get_object_or_404(Instructor, pk=idInstructor)

    if request.method == 'POST':
        form = AppointmentForm(request.POST)
        if form.is_valid():
            date = form.cleaned_data['date']
            time = request.POST.get('time')

            existing_appointment = Appointment.objects.filter(date=date, time=time, instructor=instructor).first()
            if existing_appointment:
                    return render(request, 'student_pickDataTime.html', {'form': form, 'error_message': 'Время уже занято.'})
            else:
                appointment = form.save(commit=False)
                appointment.instructor = instructor
                appointment.is_available = False
                appointment.time = time
                appointment.student = request.user
                appointment.save()
                context = {'appointments': appointment}
                return render(request, 'success_url.html', context)
    else:
        form = AppointmentForm()

        date = form['date'].value()
        existing_appointments = Appointment.objects.filter(date=date)
        occupied_times = [appointment.time for appointment in existing_appointments if not appointment.is_available]
        available_times_filtered = [t for t in available_times if t not in occupied_times]

        time_choices = [(t.strftime('%H:%M'), t.strftime('%H:%M')) for t in available_times_filtered]
        form.fields['time'] = forms.ChoiceField(choices=time_choices)

    return render(request, 'student_pickDataTime.html', {'form': form})



def get_available_times(request):
    selected_date = request.GET.get("date")
    instructor_id = request.GET.get("instructor_id")
    
    existing_appointments = Appointment.objects.filter(date=selected_date, instructor_id=instructor_id)
    occupied_times = [appointment.time.strftime('%H:%M') for appointment in existing_appointments if not appointment.is_available]

    
    available_times = ["08:00", "09:30", "11:00", "13:30", "15:00", "16:30"]
    
    available_times = [time for time in available_times if time not in occupied_times]
    
    return JsonResponse(available_times, safe=False)

def schedule(request):
    user_id = request.user.id  
    appointments = Appointment.objects.filter(student_id=user_id)
    
    appointments_data = []
    for appointment in appointments:
        appointments_data.append({
            'id': appointment.id,
            'date': appointment.date.strftime('%Y-%m-%d'),
            'time': appointment.time.strftime('%H:%M:%S'),
            'instructor': f"{appointment.instructor.second_name} {appointment.instructor.name} {appointment.instructor.surname}",
        })

    appointments_json = json.dumps(appointments_data)

    context = {'appointments_json': appointments_json, 'user_id': user_id}
    return render(request, 'schedule.html', context)

    
def teacher_dashboard(request):
    user_id = request.user.id 
    instructor = Instructor.objects.get(user_id=user_id)
    instructor_id = instructor.id
    
    appointments = Appointment.objects.filter(instructor_id=instructor_id)
    
    appointments_data = []
    for appointment in appointments:
        appointments_data.append({
            'id': appointment.id,
            'date': appointment.date.strftime('%Y-%m-%d'),
            'time': appointment.time.strftime('%H:%M:%S'),
            'student': f"{appointment.student.first_name} {appointment.student.last_name}",
            'student_id': appointment.student.id,
        })

    appointments_json = json.dumps(appointments_data)

    context = {'appointments_json': appointments_json, 'user_id': instructor_id}
    return render(request, 'teacher_dashboard.html', context)

@csrf_exempt
@require_POST
def delete_event(request):
    try:
        data = json.loads(request.body)
        event_date = data.get('event_date')
        event_time = data.get('event_time')

        if event_date is not None and event_time is not None:
            appointment = Appointment.objects.filter(date=event_date, time=event_time)
            if appointment.exists():
                appointment.delete()
                return JsonResponse({'status': 'success'})
            else:
                return JsonResponse({'status': 'error', 'message': 'Событие не найдено.'})
        else:
            return JsonResponse({'status': 'error', 'message': 'Дата и/или время события не указаны.'})

    except json.JSONDecodeError as e:
        return JsonResponse({'status': 'error', 'message': str(e)})

    

def teacher_notation(request):
    group = Group.objects.get(id=2)
    users_sorted = User.objects.filter(email_verify=True, groups=group)

    user_id_current = request.user.id 
    instructor = Instructor.objects.get(user_id=user_id_current)
    instructor_id = instructor.id


    if request.method == 'POST':
        form = AppointmentForm(request.POST)
        if form.is_valid():
            date = form.cleaned_data['date']
            time = request.POST.get('time')

            existing_appointment = Appointment.objects.filter(date=date, time=time, instructor=instructor_id).first()
            if existing_appointment:
                    return render(request, 'teacher_notation.html', {'form': form, 'error_message': 'Время уже занято.'})
            else:
                appointment = form.save(commit=False)
                appointment.instructor = instructor
                appointment.is_available = False
                appointment.time = time
                selected_user_id = request.POST.get('user_radio')
                selected_user = User.objects.get(id=selected_user_id)
                appointment.student = selected_user
                appointment.save()
                context = {'appointments': appointment}
                return render(request, 'success_url_instructor.html', context)
    else:
        form = AppointmentForm()

        date = form['date'].value()
        existing_appointments = Appointment.objects.filter(date=date)
        occupied_times = [appointment.time for appointment in existing_appointments if not appointment.is_available]
        available_times_filtered = [t for t in available_times if t not in occupied_times]

        time_choices = [(t.strftime('%H:%M'), t.strftime('%H:%M')) for t in available_times_filtered]
        form.fields['time'] = forms.ChoiceField(choices=time_choices)

    context = {'users': users_sorted, 'form': form}

    return render(request, 'teacher_notation.html', context)

def  teacher_notice(request):
    group = Group.objects.get(id=2)
    users_sorted = User.objects.filter(email_verify=True, groups=group)
    context = {'users': users_sorted}
    return render(request, 'teacher_notice.html', context)

@login_required
def send_message(request):
    if request.method == 'POST':
        recipient_ids_str = request.POST.get('recipient_ids') 
        recipient_ids = recipient_ids_str.split(',')
        message_body = request.POST.get('textMessage')  

        sender = request.user

        for recipient_id in recipient_ids:
            recipient = User.objects.get(id=recipient_id)
            message = Message.objects.create(sender=sender, recipient=recipient, body=message_body)
            message.save()
        messages.success(request, 'Сообщение успешно отправлено!')
        return redirect('teacher_notice')

def manager_dashboard(request):
    users = User.objects.filter(email_verify=0) 
    return render(request, 'manager_dashboard.html', {'users': users})

@require_POST
def confirm_users(request):
    user_ids = request.POST.getlist('user_id')  
    users = User.objects.filter(id__in=user_ids)
    print(users)

    user_ids_before_update = list(users.values_list('id', flat=True))

    users.update(email_verify=1)  

    group = Group.objects.get(id=2) 
    for user in users:
        user.groups.add(group)

    response_data = {'reload_page': True}
    return JsonResponse(response_data)



def manager_add_users_data(request):
    group = Group.objects.get(id=2)
    users = User.objects.filter(email_verify=True, groups=group)
    return render(request, 'manager_add_users_data.html', {'users': users})


def add_user_data(request, user_id):
    user = User.objects.get(id=user_id)
    user_data, created = UserData.objects.get_or_create(user=user)

    formatted_date = user_data.date_of_birth.strftime('%Y-%m-%d') if user_data.date_of_birth else None

    if request.method == 'POST':
        form = UserDataForm(request.POST, instance=user_data)
        if form.is_valid():
            form.save()
            return redirect('manager_add_users_data')
    else:
        form = UserDataForm(instance=user_data)

    return render(request, 'add_user_data.html', {'user': user, 'form': form, 'formatted_date': formatted_date})


def payments(request):
    user = request.user
    payments = Payments.objects.filter(user=user)
    total_payment = payments.aggregate(Sum('payment'))['payment__sum'] or 0
    return render(request, 'student_payments.html', {'payments': payments, 'total_payment': total_payment})

def manager_add_users_payments(request):
    group = Group.objects.get(id=2)
    users = User.objects.filter(email_verify=True, groups=group)
    return render(request, 'add_users_payments.html', {'users': users})

def add_user_payment(request, user_id):
    user = User.objects.get(id=user_id)

    if request.method == 'POST':
        form = UserPaymentsForm(request.POST)
        if form.is_valid():
            form.instance.user = user 
            form.save()
            return redirect('manager_add_users_payments')
    else:
        form = UserPaymentsForm()

    return render(request, 'add_user_payment.html', {'user': user, 'form': form})

def profile(request):
    user_id = request.user.id
    user = User.objects.get(id=user_id)
    two_weeks_ago = timezone.now() - timedelta(weeks=2)
    
    received_messages = Message.objects.filter(recipient=user, timestamp__gte=two_weeks_ago).order_by('-timestamp')

    return render(request, 'student_profile.html', {'user': user, 'received_messages': received_messages})


def generate_group_journal(request):
    if request.method == 'POST':
        group_number = request.POST.get('group_number')
        
        user_data_list = UserData.objects.filter(group_number=group_number)

        if user_data_list:
            template_path = "group_journal.docx"

            doc = Document(template_path)

            def set_font_and_size(cell):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial CYR'
                        run.font.size = Pt(10)
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        paragraph.paragraph_format.space_after = Pt(0)
            def set_font_and_size_without_center(cell):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial CYR'
                        run.font.size = Pt(10)
                        paragraph.paragraph_format.space_after = Pt(0)

            row_table_indices = [0, 1, 2, 4, 6, 8, 10, 12, 14, 16, 18, 20]
            for idx in row_table_indices:
                row_table = doc.tables[idx]
                num_needed_rows = len(user_data_list)
                num_existing_rows = len(row_table.rows)
                if (idx == 2) or (idx == 4) or (idx == 6) or (idx == 8) or (idx == 10) or (idx == 12) or (idx == 14) or (idx == 16) or (idx == 18):
                    if num_existing_rows < num_needed_rows + 3:  
                        num_rows_to_add = num_needed_rows - num_existing_rows + 3
                        for _ in range(num_rows_to_add):
                            row_table.add_row()
                if (idx == 1):
                    if num_existing_rows < num_needed_rows + 2:  
                        num_rows_to_add = num_needed_rows - num_existing_rows + 2
                        for _ in range(num_rows_to_add):
                            row_table.add_row()
                if (idx == 0) or (idx == 20):
                    if num_existing_rows < num_needed_rows + 1:  
                        num_rows_to_add = num_needed_rows - num_existing_rows + 1
                        for _ in range(num_rows_to_add):
                            row_table.add_row()
            
            for i, user_data in enumerate(user_data_list):
                user = user_data.user
                context = {
                    'last_name': user.last_name if user.last_name else '',
                    'first_name': user.first_name if user.first_name else '',
                    'surname': user.surname if user.surname else '',
                    'date_of_birth': user_data.date_of_birth.strftime('%d.%m.%Y') if user_data.date_of_birth else '',
                    'region': user_data.region if user_data.region else '',
                    'city_or_village': user_data.city_or_village if user_data.city_or_village else '',
                    'street': user_data.street if user_data.street else '',
                    'house': user_data.house if user_data.house else '',
                    'corps': user_data.corps if user_data.corps else '',
                    'apartment': user_data.apartment if user_data.apartment else '',
                }
                
                row = doc.tables[0].rows[i + 1].cells  
                row[0].text = str(i + 1)  
                row[1].text = context['last_name']  
                set_font_and_size(row[1])
                row[2].text = context['first_name']  
                set_font_and_size(row[2])
                row[3].text = context['surname']  
                set_font_and_size(row[3])
                row[4].text = context['date_of_birth']  
                set_font_and_size(row[4])

                row_table2 = doc.tables[1].rows[i + 2].cells
                row_table2[0].text = context['region']
                set_font_and_size(row_table2[0])
                row_table2[1].text = context['city_or_village'] 
                set_font_and_size(row_table2[1])   
                row_table2[2].text = context['street']
                set_font_and_size(row_table2[2])
                row_table2[3].text = context['house'] 
                set_font_and_size(row_table2[3])
                row_table2[4].text = context['corps']  
                set_font_and_size(row_table2[4])
                row_table2[5].text = context['apartment']
                set_font_and_size(row_table2[5])

                row_table3 = doc.tables[2].rows[i+3].cells
                row_table3[0].text = str(i + 1)
                set_font_and_size_without_center(row_table3[0])
                row_table3[1].text = context['last_name'] + ' '+ context['first_name'][0] + '. ' + context['surname'][0] + '. '
                set_font_and_size_without_center(row_table3[1])

                row_table4 = doc.tables[4].rows[i+3].cells
                row_table4[0].text = str(i + 1)
                set_font_and_size_without_center(row_table4[0])
                # row_table3 = doc.tables[2].rows[i+3].cells
                row_table4[1].text = context['last_name'] + ' '+ context['first_name'][0] + '. ' + context['surname'][0] + '. '
                set_font_and_size_without_center(row_table4[1])

                row_table5 = doc.tables[6].rows[i+3].cells
                row_table5[0].text = str(i + 1)
                set_font_and_size_without_center(row_table5[0])
                row_table5[1].text = context['last_name'] + ' '+ context['first_name'][0] + '. ' + context['surname'][0] + '. '
                set_font_and_size_without_center(row_table5[1])

                row_table6 = doc.tables[8].rows[i+3].cells
                row_table6[0].text = str(i + 1)
                set_font_and_size_without_center(row_table6[0])
                row_table6[1].text = context['last_name'] + ' '+ context['first_name'][0] + '. ' + context['surname'][0] + '. '
                set_font_and_size_without_center(row_table6[1])

                row_table7 = doc.tables[10].rows[i+3].cells
                row_table7[0].text = str(i + 1)
                set_font_and_size_without_center(row_table7[0])
                row_table7[1].text = context['last_name'] + ' '+ context['first_name'][0] + '. ' + context['surname'][0] + '. '
                set_font_and_size_without_center(row_table7[1])

                row_table8 = doc.tables[12].rows[i+3].cells
                row_table8[0].text = str(i + 1)
                set_font_and_size_without_center(row_table8[0])
                row_table8[1].text = context['last_name'] + ' '+ context['first_name'][0] + '. ' + context['surname'][0] + '. '
                set_font_and_size_without_center(row_table8[1])

                row_table9 = doc.tables[14].rows[i+3].cells
                row_table9[0].text = str(i + 1)
                set_font_and_size_without_center(row_table9[0])
                row_table9[1].text = context['last_name'] + ' '+ context['first_name'][0] + '. ' + context['surname'][0] + '. '
                set_font_and_size_without_center(row_table9[1])

                row_table10 = doc.tables[16].rows[i+3].cells
                row_table10[0].text = str(i + 1)
                set_font_and_size_without_center(row_table10[0])
                row_table10[1].text = context['last_name'] + ' '+ context['first_name'][0] + '. ' + context['surname'][0] + '. '
                set_font_and_size_without_center(row_table10[1])

                row_table11 = doc.tables[18].rows[i+3].cells
                row_table11[0].text = str(i + 1)
                set_font_and_size_without_center(row_table11[0])
                row_table11[1].text = context['last_name'] + ' '+ context['first_name'][0] + '. ' + context['surname'][0] + '. '
                set_font_and_size_without_center(row_table11[1])

                row_table12 = doc.tables[20].rows[i+1].cells
                row_table12[0].text = str(i + 1)
                set_font_and_size_without_center(row_table12[0])
                row_table12[1].text = context['last_name'] + ' '+ context['first_name'][0] + '. ' + context['surname'][0] + '. '
                set_font_and_size_without_center(row_table12[1])
            
            for p in doc.paragraphs:
                if "{{group_number}}" in p.text:
                    p.text = p.text.replace("{{group_number}}", group_number)
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        run.font.size = Pt(36)
                    break 

            
            output_path = "filled_group_journal.docx"
            doc.save(output_path)

            with open(output_path, 'rb') as docx_file:
                response = HttpResponse(docx_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename=filled_group_journal.docx'
                return response
                
        else:
            return HttpResponse("Пользователи с указанной группой не найдены")

    return render(request, 'manager_documents.html')

def manager_document_dogovor(request):
    # users = User.objects.filter(email_verify=1) & User.objects.filter(Group.objects.get(id=2))
    group = Group.objects.get(id=2)
    users = User.objects.filter(email_verify=True, groups=group)
    return render(request, 'manager_document_dogovor.html', {'users': users})

def generate_docx(user):
    doc = Document('dogovor_ob_obych_В.docx')

    user_data = user.userdata

    context = {
                'region': user_data.region if user_data.region else '',
                'city_or_village': user_data.city_or_village if user_data.city_or_village else '',
                'street': user_data.street if user_data.street else '',
                'house': user_data.house if user_data.house else '',
                'corps': "корпус " + user_data.corps if user_data.corps else '',
                'apartment': user_data.apartment if user_data.apartment else '',
            }
    
    for paragraph in doc.paragraphs:
        if "{{last_name}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{last_name}}", user.last_name)
        if "{{first_name}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{first_name}}", user.first_name)
        if "{{surname}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{surname}}", user.surname)
        if "{{date_of_birth}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{date_of_birth}}", str(user_data.date_of_birth.day))
        if "{{month_of_birth}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{month_of_birth}}", str(user_data.date_of_birth.month))
        if "{{year_of_birth}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{year_of_birth}}", str(user_data.date_of_birth.year))
        if "{{place_of_birth}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{place_of_birth}}", str(user_data.place_of_birth))
        if "{{passport_series}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{passport_series}}", str(user_data.passport_series))
        if "{{passport_number}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{passport_number}}", str(user_data.passport_number))
        if "{{passport_day}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{passport_day}}", str(user_data.passport_date.day))
        if "{{passport_month}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{passport_month}}", str(user_data.passport_date.month))
        if "{{passport_year}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{passport_year}}", str(user_data.passport_date.year))
        if "{{passport_year}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{passport_year}}", str(user_data.passport_date.year))
        if "{{region}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{region}}", context['region'] )
        if "{{city_or_village}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{city_or_village}}", context['city_or_village'] )
        if "{{street}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{street}}", context['street'] )
        if "{{house}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{house}}", context['house'] )
        if "{{corps}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{corps}}", context['corps'] )
        if "{{apartment}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{apartment}}", context['apartment'] )
        if "{{phone}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{phone}}", str(user.phone))
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def user_dogovor(request, user_id):
    print(user_id)
    user = get_object_or_404(User, id=user_id)
    
    docx_buffer = generate_docx(user)
    
    response = HttpResponse(docx_buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="document.docx"'
    
    return response

def generate_hour_group(request):
    if request.method == 'POST':
        group_number = request.POST.get('group_number')
        
        user_data_list = UserData.objects.filter(group_number=group_number)

        def set_font(cell):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        if user_data_list:
            template_path = "накопительная_ведомость.docx"

            doc = Document(template_path)

            for paragraph in doc.paragraphs:
                if "{{group_number}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{group_number}}", group_number)

            
            for i, user_data in enumerate(user_data_list):
                user = user_data.user
                
                context = {
                    'last_name': user.last_name if user.last_name else '',
                    'first_name': user.first_name if user.first_name else '',
                    'surname': user.surname if user.surname else '',
                }

                row = doc.tables[0].rows[i+1].cells  
                row[1].text = context['last_name'] + ' '+ context['first_name'][0] + '. ' + context['surname'][0] + '. '

                student_appointments = Appointment.objects.filter(student=user)
                appointment_month = [str(appointment.date.month) for appointment in student_appointments]
                appointment_day = [str(appointment.date.day) for appointment in student_appointments]
                
                appointment_dates_str = ', '.join(appointment_month)
                
                if appointment_dates_str:
                    for j, date in enumerate(appointment_month, start=2):
                        if j < 22: 
                            if date < '10':
                                row[j].text = appointment_day[j-2] + '.' + '0' + date
                                set_font(row[j])
                            else:
                                row[j].text = appointment_day[j-2] + '.' + date 
                                set_font(row[j])
                        if j >= 22 and j < 30: 
                            if date < '10':
                                row_table2 = doc.tables[1].rows[i+1].cells
                                row_table2[j-20].text = appointment_day[j-2] + '.' + '0' + date
                                set_font(row_table2[j-20])
                            else:
                                row_table2 = doc.tables[1].rows[i+1].cells
                                row_table2[j-20].text = appointment_day[j-2] + '.' + date 
                                set_font(row_table2[j-20])
                            



                row_table2 = doc.tables[1].rows[i+1].cells  
                row_table2[1].text = context['last_name'] + ' '+ context['first_name'][0] + '. ' + context['surname'][0] + '. '


            
            output_path = "накопительная_ведомостьfill.docx"
            doc.save(output_path)

            with open(output_path, 'rb') as docx_file:
                response = HttpResponse(docx_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                new_filename = "накопительная_ведомость№"+group_number+".docx"
                response['Content-Disposition'] = 'attachment; filename*=UTF-8\'\'' + quote(new_filename)
                return response
                
        else:
            return HttpResponse("Пользователи с указанной группой не найдены")

    return render(request, 'manager_document_group.html')


@require_POST
def kniga_vojden_users(request):
    user_id = request.POST.get('user_id')  
    user = User.objects.get(id=user_id)

    instructor_id = request.POST.get('instructor_id')  
    instructor = Instructor.objects.get(id__in=instructor_id)

    print(user)
    print(instructor)

    if instructor_id:

        if (instructor_id == '1') : 
            template_path = "книжка вождения Орлова.docx"
            
        if (instructor_id == '2') : 
            template_path = "книжка вождения Кузьмичев.docx"

        if (instructor_id == '3') : 
            template_path = "книжка вождения Русаков.docx"

        if (instructor_id == '4') : 
            template_path = "книжка вождения Алиференко.docx"
        
        doc = Document(template_path)

        context = {
                    'last_name': user.last_name if user.last_name else '',
                    'first_name': user.first_name if user.first_name else '',
                    'surname': user.surname if user.surname else '',
                }
        
        for paragraph in doc.paragraphs:
            if "{{last_name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{last_name}}", context['last_name'])                  
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'  
                    run.font.size = Pt(16)      
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
            if "{{first_name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{first_name}}", context['first_name'])                
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'  
                    run.font.size = Pt(16) 
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if "{{surname}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{surname}}", context['surname'])                
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'  
                    run.font.size = Pt(16) 
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

       
        student_appointments = Appointment.objects.filter(student=user)
        appointment_dates = [f"{appointment.date.day}.{appointment.date.month}" for appointment in student_appointments]
        print(appointment_dates)

        idx = 0

        table_count = 0

        for element in doc.element.body:
            if element.tag.endswith('tbl'):
                table_count += 1
                if table_count == 5:
                    for row_idx, row in enumerate(element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')):
                        for cell_idx, cell in enumerate(row.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')):
                            for paragraph in cell.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                                for text_element in paragraph.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                                    if "{{date}}" in text_element.text:
                                        date_str = appointment_dates[0]  
                                        date_parts = date_str.split('.')  
                                        day = int(date_parts[0])
                                        month = int(date_parts[1])
                                        if month < 10:
                                            text_element.text = text_element.text.replace("{{date}}", str(day) + '.0' + str(month))
                                        else:
                                            text_element.text = text_element.text.replace("{{date}}", appointment_dates[0])
                                            # set_font(row[j])
                                        # text_element.text = text_element.text.replace("{{date}}", appointment_dates[0])
                                        appointment_dates = appointment_dates[1:]
                                        
        
        output_path = "книжка вождения " + user.last_name + user.first_name + ".docx"
        doc.save(output_path)

        with open(output_path, 'rb') as docx_file:
            response = HttpResponse(docx_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            new_filename = "книжка вождения Орлова " + user.last_name + ".docx"
            response['Content-Disposition'] = 'attachment; filename*=UTF-8\'\'' + quote(new_filename)
            return response
        
    return render(request, 'manager_document_kniga_vojden.html')

def kniga_vojden(request):
    group = Group.objects.get(id=2)
    users = User.objects.filter(email_verify=True, groups=group) 
    instructors = Instructor.objects.all()

    context = {'users': users, 'instructors': instructors}
    
    return render(request, 'manager_document_kniga_vojden.html', context)

def document_instructor(request):
    instructors = Instructor.objects.all()

    context = {'instructors': instructors}
    
    return render(request, 'manager_document_instructor.html', context)


@require_POST
def instructor_lessons(request):
    month_number = request.POST.get('month_number') 

    instructor_id = request.POST.get('instructor_id')  
    instructor = Instructor.objects.get(id__in=instructor_id)

    if (instructor_id and month_number):
        appointments = Appointment.objects.filter(
            instructor_id=instructor_id,
            date__month=month_number
        ).values('date').annotate(count=Count('id'))

        appointments_count_dict = {entry['date']: entry['count'] for entry in appointments}
        num_days_in_month = calendar.monthrange(datetime.now().year, int(month_number))[1]

        print(month_number)
        print(instructor)

        template_path = "Табель учета рабочего времени мастера ПОВА.docx"
        doc = Document(template_path)

        context = {
                    'second_name': instructor.second_name if instructor.second_name else '',
                    'name': instructor.name if instructor.name else '',
                    'surname': instructor.surname if instructor.surname else '',
                }
        
        for paragraph in doc.paragraphs:
            if "{{second_name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{second_name}}", context['second_name'])                  
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'  
                    run.font.size = Pt(16)      
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
            if "{{name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{name}}", context['name'])                
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'  
                    run.font.size = Pt(16) 
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if "{{surname}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{surname}}", context['surname'])                
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'  
                    run.font.size = Pt(16) 
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for i in range(num_days_in_month):
                # row[1].text = instructor.second_name
                date_for_current_day = datetime(datetime.now().year, int(month_number), i+1).date()
                count_for_current_day = appointments_count_dict.get(date_for_current_day, 0)
                row = doc.tables[0].rows[i+3].cells
                row[1].text = str(count_for_current_day)
                # row[2].text = str(count_for_current_day)
                print(date_for_current_day)
                print(count_for_current_day)

        
        output_path = "Отчет " + instructor.second_name + ".docx"
        doc.save(output_path)

        with open(output_path, 'rb') as docx_file:
            response = HttpResponse(docx_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            new_filename = "Отчет " + instructor.second_name + ".docx"
            response['Content-Disposition'] = 'attachment; filename*=UTF-8\'\'' + quote(new_filename)
            return response
        
    return render(request, 'manager_document_instructor.html')


# def send_reminder_emails():
#     # Получаем текущую дату
#     current_date = timezone.now().date()
    
#     # Получаем список всех занятий, которые состоятся завтра
#     tomorrow_appointments = Appointment.objects.filter(date=current_date + timedelta(days=1))
    
#     for appointment in tomorrow_appointments:
#         # Отправляем уведомление студенту
#         send_mail(
#             'Напоминание о занятии',
#             f'Завтра, {appointment.date}, в {appointment.time} у вас запланировано занятие с инструктором {appointment.instructor}.',
#             'your_email@yandex.ru',
#             [appointment.student.email],
#             fail_silently=False,
#         )